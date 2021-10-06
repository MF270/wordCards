"""Makes your word cards for you!!"""
# coding: utf8
import csv
import sys
import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_BREAK
from docx.shared import Inches,Pt

def genCards(inputPath:str="",outputPath:str="/",userName:str="",userLesson:str=""):
    """Makes and outputs cards given a csv and file output"""
    #Sets basic formatting for the whole doc like setup + header
    if not inputPath or not outputPath:
        sys.exit()
    terms = []
    document = Document()
    sections = document.sections
    sections[0].orientation = WD_ORIENT.LANDSCAPE
    sections[0].page_width = Inches(6)
    sections[0].page_height = Inches(4)
    sections[0].header.paragraphs[0].text = f"{userName} \n {userLesson}"
    #Iterates through each sub-list in terms
    with open(inputPath,encoding='utf-8') as csv_file:
        csv_reader = csv.reader(csv_file,delimiter=',')
        for row in csv_reader:
            character = [row[0],row[1],row[2]]
            terms.append(character)
    for term in terms:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(term[0])
        run.font.name = 'Microsoft JhengHei UI'
        if len(term[0]) <=2:
            run.font.size = Pt(96)
        elif len(term[0]) == 3:
            run.font.size = Pt(72)
        else:
            run.font.size = Pt(56)
        #Characters should be big
        run.add_break(WD_BREAK.PAGE)
        run = paragraph.add_run(term[1] + "\n")
        run.font.name = 'Calibri'
        if len(term[1])<12:
            run.font.size = Pt(48)
        else:
            run.font.size = Pt(36)
        #Pinyin pretty much always fits in size 48
        run = paragraph.add_run(term[2])
        run.font.name = 'Calibri'
        if len(term[2]) > 18:
            run.font.size = Pt(18)
        elif len(term[2]) > 13:
            run.font.size = Pt(24)
        elif len(term[2]) > 8:
            run.font.size = Pt(32)
        else:
            run.font.size = Pt(48)
        #Definitions can get long, so they need to be smaller

    document.save(f'{outputPath}/{userLesson}cards.docx')
    #Bam, saved as a docx
    print("Done!")

def displayDir(var):
    """Makes the output 'browse' button work"""
    path = filedialog.askdirectory()
    var.set(path.name)

def openInput(var):
    """Makes the input 'browse' button work"""
    words = filedialog.askopenfile(mode="r",initialdir="/")
    var.set(words.name)
    words.close()#This is almost aggressively bad code, but it's (hopefully) functional!

#Initializing window object w/title
root = tk.Tk()
root.title("Milo's Mandarin Word Card Generator")
root.geometry("750x300")

#Handles name (this only affects the header)
name = tk.StringVar()
name_label = tk.Label(root,text='Name:',font=("Helvetica",12))
name_label.config()
name_label.grid(row=0,column=0,padx=3,pady=3)
name_ent = tk.Entry(root,textvariable=name)
name_ent.grid(row=0,column=1,padx=3,pady=3,ipadx=3,ipady=3,sticky="ew")

info = "Please fill all boxes. Input file should be a spreadsheet saved as a .csv"
info2 = "Each row should read [character][pinyin][definition]"
infopanel = tk.Label(root,text=info,font=("Helvetica",10))
infopanel.grid(row=0,column=2,padx=3,pady=3,ipadx=3,ipady=3,sticky="ew")
infopanel2 = tk.Label(root,text=info2,font=("Helvetica",10))
infopanel2.grid(row=1,column=2,padx=3,pady=3,ipadx=3,ipady=3,sticky="ew")

warning = "Please make use of the browse buttons"
#Handles the lesson (affects header and document name)
lesson = tk.StringVar()
lesson_label = tk.Label(root,text='Lesson:',font=("Helvetica",12))
lesson_label.grid(row=1,column=0,padx=3,pady=3)
lesson_ent = tk.Entry(root,textvariable=lesson)
lesson_ent.grid(row=1,column=1,padx=3,pady=3,ipadx=3,ipady=3,sticky="ew")

#Handles the input path
inputpath = tk.StringVar()
input_label = tk.Label(root,text='Input File (.csv):',font=("Helvetica",12))
input_label.grid(row=2,column=0,padx=3,pady=3)
input_ent = tk.Entry(root,textvariable=input)
input_ent.grid(row=2,column=1,padx=3,pady=3,ipadx=3,ipady=3,sticky="ew")
browseOutput = tk.Button(root,text="Browse",command= lambda:
openInput(inputpath),font=("Helvetica",12))
browseOutput.grid(row=2,column=2,padx=3,pady=3,sticky="ew")

#Handles the output path, a little more complex
output = tk.StringVar()
output_label = tk.Label(root,text='Output Folder:',font=("Helvetica",12))
output_label.grid(row=3,column=0,padx=3,pady=3)
output_ent = tk.Entry(root,textvariable=output)
output_ent.grid(row=3,column=1,padx=3,pady=3,ipadx=3,ipady=3,sticky="ew")
browseOutput = tk.Button(root,text="Browse",command=
lambda: displayDir(output),font=("Helvetica",12))
browseOutput.grid(row=3,column=2,padx=3,pady=3,sticky="ew")

#Start button, when called, gets inputs from all the text boxes.
#Do not leave the input or output boxes blank. That might get screwy fast.
start = tk.Button(root,text='Start',command =
lambda : genCards(inputpath.get(),output.get(),name.get(),lesson.get()))
start.grid(row=4,column=1,padx=3,pady=3,ipadx=3,ipady=3,sticky="ew")

#Resizes the rows/columns so it looks half-decent
for i in range(4):
    root.columnconfigure(i, weight=1, minsize=50)
    root.rowconfigure(i, weight=1, minsize=50)

root.mainloop()
